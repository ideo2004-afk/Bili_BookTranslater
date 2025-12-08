import os
import sys
import pickle
import time
from pathlib import Path
from copy import deepcopy

import docx
from book_maker.utils import prompt_config_to_kwargs, global_state, num_tokens_from_text
from .base_loader import BaseBookLoader
from .helper import shorter_result_link


from .accumulation_mixin import AccumulationMixin

class DOCXBookLoader(BaseBookLoader, AccumulationMixin):
    def __init__(
        self,
        docx_name,
        model,
        key,
        resume,
        language,
        model_api_base=None,
        is_test=False,
        test_num=5,
        prompt_config=None,
        single_translate=False,
        context_flag=False,
        context_paragraph_limit=0,
        temperature=1.0,
        source_lang="auto",
        parallel_workers=1,
        glossary_path=None,
        accumulated_num=1,
    ) -> None:
        self.docx_name = docx_name
        self.accumulated_num = accumulated_num
        self.translate_model = model(
            key,
            language,
            api_base=model_api_base,
            temperature=temperature,
            source_lang=source_lang,
            glossary_path=glossary_path,
            **prompt_config_to_kwargs(prompt_config),
        )
        self.is_test = is_test
        self.p_to_save = []
        self.test_num = test_num
        self.single_translate = single_translate
        self.context_flag = context_flag
        self.parallel_workers = max(1, parallel_workers)

        try:
            self.document = docx.Document(docx_name)
        except Exception as e:
            raise Exception("can not load file") from e

        self.resume = resume
        self.bin_path = f"{Path(docx_name).parent}/.{Path(docx_name).stem}.temp.bin"
        if self.resume:
            self.load_state()

    @staticmethod
    def _is_special_text(text):
        return text.strip() == "" or text.isdigit()

    def _make_new_book(self, book):
        pass

    def _save_temp_book(self):
        pass

    def estimate(self):
        print("Calculating estimate...")
        total_tokens = 0
        for paragraph in self.document.paragraphs:
            if self._is_special_text(paragraph.text):
                continue
            # Simple token estimation
            total_tokens += len(paragraph.text) // 4
        
        print(f"Total estimated tokens: {total_tokens}")

    def make_bilingual_book(self):
        index = 0
        p_to_save_len = len(self.p_to_save)

        try:
            
            if self.accumulated_num > 1:
                # Use accumulation logic
                p_list = [p for p in self.document.paragraphs if not self._is_special_text(p.text)]
                self.translate_paragraphs_acc(p_list, self.accumulated_num, index, p_to_save_len)
            else:
                # Original line-by-line logic
                for paragraph in self.document.paragraphs:
                    if global_state.is_cancelled:
                        raise KeyboardInterrupt("Cancelled by user")

                    if self._is_special_text(paragraph.text):
                        continue

                    if not self.resume or index // 1 >= p_to_save_len:
                        try:
                            temp = self.translate_model.translate(paragraph.text)
                        except Exception as e:
                            print(f"Error during translation: {e}")
                            raise e

                        self.p_to_save.append(temp)
                    else:
                        temp = self.p_to_save[index]

                    self._update_paragraph(paragraph, temp)

                    index += 1
                    if index % 20 == 0:
                        self._save_progress()

                    if self.is_test and index > self.test_num:
                        break

            self.save_file(
                f"{Path(self.docx_name).parent}/{Path(self.docx_name).stem}_bili.docx"
            )

        except (KeyboardInterrupt, Exception) as e:
            print(e)
            print("you can resume it next time")
            self._save_progress()
            sys.exit(1)
            
        # Print Performance Summary
        if hasattr(self.translate_model, 'total_tokens') and hasattr(self.translate_model, 'total_time'):
            total_tokens = self.translate_model.total_tokens
            total_time = self.translate_model.total_time
            avg_speed = total_tokens / total_time if total_time > 0 else 0
            
            print("\n" + "="*50)
            print("📊 Translation Performance Summary")
            print("="*50)
            print(f"Model: {self.translate_model.model}")
            print(f"Total Tokens Processed: {total_tokens:,}")
            print(f"Total Translation Time: {total_time:.2f}s")
            print(f"Average Speed: {avg_speed:.2f} tokens/s")
            print("="*50 + "\n")
            
            # Optional: Save to file
            try:
                os.makedirs("log", exist_ok=True)
                with open("log/translation_stats.txt", "a", encoding="utf-8") as f:
                    f.write(f"\n--- {time.strftime('%Y-%m-%d %H:%M:%S')} ---\n")
                    f.write(f"Book: {self.docx_name}\n")
                    f.write(f"Model: {self.translate_model.model}\n")
                    f.write(f"Total Tokens: {total_tokens}\n")
                    f.write(f"Total Time: {total_time:.2f}s\n")
                    f.write(f"Avg Speed: {avg_speed:.2f} t/s\n")
            except Exception as e:
                print(f"Failed to save stats: {e}")

    def _update_paragraph(self, paragraph, temp):
        if self.single_translate:
            paragraph.text = temp
        else:
            # Insert translation after the original paragraph
            # For simplicity in this first version, we just append to the text
            paragraph.text = f"{paragraph.text}\n{temp}"

    def _save_progress(self):
        try:
            with open(self.bin_path, "wb") as f:
                pickle.dump(self.p_to_save, f)
        except Exception as e:
            raise Exception("can not save resume file") from e

    def load_state(self):
        try:
            with open(self.bin_path, "rb") as f:
                self.p_to_save = pickle.load(f)
        except FileNotFoundError:
            print("Resume file not found, starting from beginning.")
            self.p_to_save = []
        except Exception as e:
            print(f"Error loading resume file: {e}, starting from beginning.")
            self.p_to_save = []

    def save_file(self, book_path):
        try:
            self.document.save(book_path)
        except Exception as e:
            raise Exception("can not save file") from e
